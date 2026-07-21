# -*- coding: utf-8 -*-
"""
Generador de la memoria técnica del proyecto — COMO_SE_DESARROLLO_BES_DESIGNER.docx
====================================================================================
Produce el documento Word (python-docx) en la raíz del proyecto. Se puede
regenerar cuando el proyecto evolucione: editar el CONTENIDO abajo y correr

    python generar_memoria_tecnica.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).parent.parent
OUT = ROOT / "COMO_SE_DESARROLLO_BES_DESIGNER.docx"
AZUL = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()

# ---------------------------------------------------------------- estilos
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(11)
for i, sz in [(1, 17), (2, 14), (3, 12)]:
    h = doc.styles[f"Heading {i}"]
    h.font.name = "Calibri"; h.font.size = Pt(sz); h.font.color.rgb = AZUL
if "Diagrama" not in [s.name for s in doc.styles]:
    d = doc.styles.add_style("Diagrama", WD_STYLE_TYPE.PARAGRAPH)
    d.font.name = "Consolas"; d.font.size = Pt(9)
if "Ecuacion" not in [s.name for s in doc.styles]:
    e = doc.styles.add_style("Ecuacion", WD_STYLE_TYPE.PARAGRAPH)
    e.font.name = "Cambria Math"; e.font.size = Pt(12)

# ------------------------------------------------------------- helpers
def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)

def p(t, bold=False, italic=False):
    par = doc.add_paragraph()
    r = par.add_run(t); r.bold = bold; r.italic = italic
    return par

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

def eq(t, nota=None):
    par = doc.add_paragraph(style="Ecuacion")
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run(t)
    if nota:
        n = doc.add_paragraph(); n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = n.add_run(nota); r.italic = True; r.font.size = Pt(9)

def diagrama(lineas):
    for ln in lineas:
        par = doc.add_paragraph(ln, style="Diagrama")
        par.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()

def tabla(encabezados, filas, anchos=None):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    for i, e in enumerate(encabezados):
        c = t.rows[0].cells[i]
        c.text = e
        for r in c.paragraphs[0].runs: r.bold = True
    for fila in filas:
        cells = t.add_row().cells
        for i, v in enumerate(fila):
            cells[i].text = str(v)
    doc.add_paragraph()

def archivo(ruta, descripcion):
    par = doc.add_paragraph(style="List Bullet")
    r = par.add_run(ruta); r.font.name = "Consolas"; r.font.size = Pt(9.5); r.bold = True
    par.add_run(" — " + descripcion)

# ============================================================ PORTADA
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("\n\n\nBES DESIGNER"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = AZUL
tp2 = doc.add_paragraph(); tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp2.add_run("Memoria técnica del desarrollo\n").font.size = Pt(18)
tp3 = doc.add_paragraph(); tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp3.add_run("Cómo se construyó, cómo funciona, cómo mantenerlo,\n"
            "cómo ampliarlo y cómo defenderlo\n\n").italic = True
tp4 = doc.add_paragraph(); tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp4.add_run("Diseño automatizado de sistemas de Bombeo Electrosumergible (BES/ESP)\n"
            "según Kermit E. Brown, The Technology of Artificial Lift Methods, Vol. 2B, Cap. 4.5\n\n"
            "Pablo Tejerina — Tesis de grado, Ingeniería en Petróleo")
doc.add_page_break()

# ============================================================ 1. INTRODUCCIÓN
h1("1. Introducción")
p("Este documento explica cómo se desarrolló BES Designer, para que un ingeniero "
  "de producción que nunca participó del proyecto —y que no sabe programar— pueda "
  "entender cómo se construyó, cómo funciona por dentro, cómo se mantiene, cómo se "
  "amplía con nuevos fabricantes y cómo se defiende académicamente. Cada capítulo "
  "empieza con el concepto de ingeniería y recién después muestra cómo lo "
  "implementa el software, nombrando los archivos y funciones reales del proyecto.")
p("La regla de lectura es simple: el software no inventa ingeniería. Todo lo que "
  "calcula está en la bibliografía (Brown, correlaciones PVT clásicas, normas API); "
  "el programa solo lo automatiza, lo hace repetible y lo documenta. Cuando en este "
  "documento aparece el nombre de una función —por ejemplo calculate_tdh()— debe "
  "leerse como el nombre del casillero donde vive un cálculo que el lector ya "
  "conoce de la cátedra.")
h3("Vocabulario mínimo de programación (las 10 palabras que hacen falta)")
tabla(["Término", "Qué significa, en criollo"], [
    ("Función", "Una 'receta' con nombre: recibe datos, hace un cálculo, devuelve un resultado. Como una celda de Excel con fórmula, pero con nombre propio."),
    ("Módulo / archivo .py", "Un archivo que agrupa funciones afines. Como un capítulo de carpeta: 'todo lo de TDH está en tdh.py'."),
    ("Paquete / carpeta", "Un grupo de módulos. core/ es el paquete de cálculos; ui/ el de pantallas."),
    ("Clase / objeto", "Un 'formulario tipo' y sus copias llenas. La clase Reservoir define qué campos tiene un yacimiento; cada pozo cargado es un objeto Reservoir con sus valores."),
    ("Diccionario", "Una lista de pares nombre→valor, como una fila de Excel con encabezados."),
    ("Importación", "Un módulo 'pide prestadas' funciones de otro. app.py importa el recomendador para usarlo."),
    ("JSON", "Un formato de archivo de texto para guardar datos estructurados. Era el formato viejo de los catálogos."),
    ("Base de datos", "Datos organizados en tablas relacionadas entre sí, guardados fuera del programa."),
    ("Interpolación", "Estimar un valor entre dos puntos conocidos de una curva (regla de tres sobre el tramo)."),
    ("API / interfaz", "El 'mostrador' de un módulo: qué se le puede pedir y qué entrega, sin importar cómo lo resuelve por dentro."),
])

# ============================================================ 2. OBJETIVOS
h1("2. Objetivos del sistema")
p("Objetivo general: automatizar el diseño completo de un sistema BES —desde los "
  "datos del pozo hasta el equipamiento recomendado— siguiendo la metodología de "
  "Brown y la enseñada en la cátedra de Producción, con resultados trazables y "
  "verificables.")
bullets([
    "Calcular el comportamiento del pozo (IPR) y las propiedades del fluido (PVT) a cualquier presión y temperatura.",
    "Calcular la presión en la admisión de la bomba (PIP) y la Altura Dinámica Total (TDH).",
    "Seleccionar bomba, número de etapas, motor, cable, sello, separador de gas, sensor y transformador desde catálogos reales.",
    "Comparar alternativas con criterios explícitos y justificar la recomendación.",
    "Analizar sensibilidad (qué pasa si cambia el caudal, el corte de agua, la presión).",
    "Generar reportes (Excel) y mantener toda la información de ingeniería en una base de datos externa al código, ampliable sin programar.",
])

# ============================================================ 3. HISTORIA
h1("3. Historia y evolución del proyecto")
p("El proyecto evolucionó en fases, y cada fase dejó una verificación automática "
  "que sigue vigente. Esta tabla es el resumen ejecutivo de la historia:")
tabla(["Fase", "Qué se hizo", "Evidencia de que funciona"], [
    ("1. Motor de cálculo", "Implementación de IPR, PVT, gradiente multifásico, TDH, etapas, potencia, según Brown §4.5.", "Los ejemplos resueltos del libro (1A, 2A, 2B, 3A, 3B) reproducidos y automatizados como tests."),
    ("2. Selección eléctrica y de gas", "Motor, cable con corrección por temperatura, sello por empuje axial, separador de gas, transformador.", "Tests unitarios por módulo (test_electrical, test_gas_handling)."),
    ("3. Recomendador e interfaz", "Puntuación de alternativas y aplicación visual Streamlit con reportes.", "Tests de integración extremo a extremo."),
    ("4. Base de datos (este trabajo)", "Migración JSON→Excel, normalización, auditoría externa, ERD, importación del catálogo Alkhorayef con digitalización de curvas.", "check_integrity.py + verify_database_v31.py + los 663 tests en verde con la base nueva."),
    ("5. Futuro", "SQLite, verificaciones MaxP/temperatura de motor/refrigeración, VSD con leyes de afinidad, pozos reales.", "La estructura ya está creada (tablas plantilla, columnas nuevas)."),
])
p("Decisión de método que atraviesa toda la historia: nunca se reescribió el "
  "proyecto. Cada mejora fue incremental, en carpetas separadas, y solo se integró "
  "al programa principal después de verificarla. La carpeta database_migration/ es "
  "el registro completo de la fase 4, con sus documentos de diseño y auditoría.")

# ============================================================ 4. DIGITALIZACIÓN
h1("4. Digitalización de catálogos y manuales")
p("Concepto: un catálogo en PDF es información para humanos; una base de datos es "
  "información para consultar. La digitalización convierte lo primero en lo "
  "segundo, sin perder la referencia al documento original.")
h3("4.1 Fichas técnicas (texto)")
p("El catálogo Alkhorayef 2019 publica una ficha por bomba (rango operativo, BEP, "
  "diámetro, límites de presión del housing, diámetro y límites del eje, tabla de "
  "housings con longitudes y pesos). Un programa lector (import_alkhorayef.py) "
  "recorre el PDF, reconoce cada ficha por sus rótulos ('Optimum operating range', "
  "'BEP flow rate'...) y copia los valores a las tablas de la base. Las 37 bombas "
  "SPECTRUM entraron así, con todos los campos que pide la metodología de la "
  "cátedra (presión máxima de housing, límites del eje, casing mínimo).")
h3("4.2 Curvas de rendimiento (gráficos)")
p("Las curvas Head–Power–Efficiency vienen como gráficos. El extractor "
  "(extract_curves_alkhorayef.py) las convierte en números en cuatro pasos: "
  "(1) extrae la imagen del gráfico del PDF; (2) lee los números de los ejes con "
  "reconocimiento óptico de caracteres (OCR) y construye la escala píxel→unidad "
  "con un ajuste robusto que tolera lecturas erróneas; (3) separa cada curva por "
  "su color (azul=altura, rojo=potencia, verde=eficiencia) y toma ~10 puntos por "
  "curva; (4) somete cada bomba a un control de calidad físico antes de aceptarla.")
p("El control de calidad es ingenieril, no informático:", bold=True)
bullets([
    "El pico de la curva de eficiencia debe caer en el caudal BEP que declara la ficha.",
    "La identidad hidráulica η = Q·H·γ / (135.773 · HP) debe cerrar punto a punto (tolerancia coherente con API: ±5% altura, ±8% potencia).",
    "La altura debe ser decreciente con el caudal (comportamiento físico de una bomba centrífuga).",
])
p("Resultado: 35 de 37 bombas pasaron automáticamente; 2 se leyeron visualmente "
  "(ejes ilegibles para el OCR) y 1 destapó una inconsistencia del propio catálogo "
  "(el pico real de su curva no coincide con el BEP declarado en su ficha) — quedó "
  "cargada con la discrepancia documentada. Además se registró el shut-off head "
  "(altura a caudal cero) de cada bomba, dato necesario para la verificación de "
  "presión máxima sobre el housing.")
p("Hallazgo documentado: la leyenda de los gráficos del catálogo tiene los colores "
  "intercambiados respecto a los ejes. Se resolvió por física (la curva que "
  "arranca alta y cae es la altura; la que tiene máximo en el BEP es la "
  "eficiencia). Es un ejemplo de por qué la digitalización requiere criterio de "
  "ingeniero y no solo software.", italic=True)

# ============================================================ 5. BASE DE DATOS
h1("5. Construcción de la base de datos")
p("Concepto: separar datos de método. Las fórmulas no cambian con el fabricante; "
  "los catálogos sí. Por eso los catálogos viven fuera del programa, en archivos "
  "Excel donde una hoja es una tabla, con claves que las relacionan — exactamente "
  "la estructura que tendrá la base SQLite futura. Migrar de Excel a SQLite será "
  "cambiar el 'archivero', no el contenido ni el programa.")
h3("5.1 Las tres bases")
tabla(["Base", "Archivos", "Contenido"], [
    ("Catálogos de equipos", "11 archivos en database_migration/data_excel/", "8 fabricantes; 60 bombas con 621 puntos de curva y housings reales; 50 motores; 19 cables con caída de tensión física; 24 sellos; 12 separadores; 4 sensores; transformadores; plantillas de VSD y switchboards."),
    ("Pozos y casos", "well_examples.xlsx, real_wells.xlsx", "Los 6 casos de validación de Brown (con resultados esperados) y la plantilla para pozos reales con instalaciones (tándems incluidos) y muestras PVT."),
    ("Bibliografía", "data_sources.xlsx", "49 fuentes con identificador corto (BROWN-01, ALKH-02...). Cada fila de cada tabla cita su fuente: ningún número queda sin referencia."),
])
h3("5.2 Decisiones de diseño defendibles")
bullets([
    "Excel primero, SQLite después: un ingeniero puede abrir, auditar y corregir Excel sin programar; la estructura lógica ya es la de SQLite.",
    "Normalización (3ª forma normal): cada dato vive en un solo lugar. La auditoría encontró y corrigió una violación real: la caída de tensión estaba repetida por cable cuando físicamente depende del conductor y el calibre.",
    "Claves legibles (Reda_400_D-40) en vez de números: las tablas se leen sin cruzarlas.",
    "Integridad verificable: check_integrity.py revisa claves únicas, referencias sin destino y rangos físicos (0 < η < 1, min < BEP < max). Detectó, por ejemplo, una serie de motor citada por sellos que no existía en el catálogo de motores.",
])
p("El plano completo de la base (tablas, claves, relaciones 1:1, 1:N y N:M) está "
  "en el diagrama entidad-relación:")
erd = ROOT / "database_migration" / "erd.png"
if erd.exists():
    doc.add_picture(str(erd), width=Inches(6.3))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Figura 1 — Diagrama entidad-relación de la base de datos (esquema v3)")
    cr.italic = True; cr.font.size = Pt(9)

# ============================================================ 6. ARQUITECTURA
h1("6. Arquitectura del software")
p("La aplicación está organizada como una empresa con departamentos: cada carpeta "
  "tiene una responsabilidad única y se comunican por 'mostradores' bien definidos. "
  "Si un departamento cambia por dentro, los demás no se enteran.")
tabla(["Carpeta", "Departamento (analogía)", "Responsabilidad"], [
    ("app.py", "Gerencia / mesa de entradas", "Orquesta todo: recibe al usuario, llama a los departamentos y muestra resultados."),
    ("core/", "Oficina de ingeniería", "Todos los cálculos: IPR, PVT, gradiente multifásico, TDH, etapas, eléctricos, gas."),
    ("catalogs/ + database_migration/", "Archivo técnico", "La base de datos y su 'ventanilla' de lectura (CatalogManager / ExcelCatalogManager)."),
    ("recommender/", "Comité de selección", "Corre el diseño con cada bomba candidata, puntúa y arma el ranking con su justificación."),
    ("ui/", "Ventanilla de atención", "Las pantallas: formularios de carga, resultados, comparación, sensibilidad, gráficos."),
    ("reports/", "Secretaría técnica", "Genera el reporte Excel descargable."),
    ("tests/", "Control de calidad", "Los 663 ensayos automáticos que validan todo contra los ejemplos de Brown."),
    ("data/", "Biblioteca de casos", "Los pozos de ejemplo del libro (hoy también en well_examples.xlsx)."),
])
p("Flujo de un diseño, de punta a punta:")
diagrama([
    "  Datos del pozo (usuario, ui/forms.py)",
    "        │",
    "        ▼",
    "  IPR — capacidad de aporte del pozo        core/ipr.py",
    "        │",
    "        ▼",
    "  PVT — propiedades del fluido              core/pvt.py",
    "        │",
    "        ▼",
    "  Gradiente multifásico → PIP               core/multiphase.py",
    "        │",
    "        ▼",
    "  TDH — altura total a vencer               core/tdh.py",
    "        │",
    "        ▼",
    "  Bomba + etapas (por candidata)            core/pump_design.py",
    "        │",
    "        ▼",
    "  Motor → sello → cable → transformador     core/electrical.py",
    "        │",
    "        ▼",
    "  Manejo de gas (si hay gas libre)          core/gas_handling.py",
    "        │",
    "        ▼",
    "  Puntuación y ranking                      recommender/",
    "        │",
    "        ▼",
    "  Pantallas y reporte Excel                 ui/  +  reports/",
])
p("Principio clave (inyección de dependencias, en criollo): los cálculos no saben "
  "de dónde salen los datos. El catálogo se crea UNA sola vez en app.py "
  "(función _load_catalog) y se les pasa a los módulos como un argumento más. Por "
  "eso cambiar JSON→Excel fue tocar una sola función, y los 663 tests lo "
  "confirmaron.")

# ============================================================ 7. MÓDULOS
h1("7. Explicación de cada módulo")
p("Para cada módulo: por qué existe, qué recibe, qué calcula, qué entrega y qué "
  "pasaría si no existiera.")

def modulo(nombre, porque, recibe, calcula, entrega, sino, funciones):
    h3(nombre)
    tabla(["Pregunta", "Respuesta"], [
        ("¿Por qué existe?", porque),
        ("¿Qué recibe?", recibe),
        ("¿Qué calcula?", calcula),
        ("¿Qué entrega?", entrega),
        ("¿Si no existiera?", sino),
        ("Funciones principales", funciones),
    ])

modulo("core/models.py — los formularios tipo",
    "Define QUÉ ES un pozo para el programa: los 'formularios' Reservoir, Fluid, WellGeometry, SurfaceConditions y DesignObjectives, con validación automática (rechaza presiones negativas, cortes de agua > 100%, frecuencias que no sean 50/60 Hz).",
    "Los valores que carga el usuario.",
    "Validaciones de coherencia física al crear cada objeto.",
    "Objetos validados que consumen todos los demás módulos, y DesignResult: el 'expediente' final del diseño.",
    "Cada módulo inventaría su propio formato y los errores de carga llegarían hasta los resultados sin aviso.",
    "Reservoir, Fluid, WellGeometry, SurfaceConditions, DesignObjectives, PumpCurve, DesignResult")

modulo("core/ipr.py — el aporte del pozo",
    "Antes de bombear hay que saber cuánto PUEDE dar el pozo a cada presión de fondo. Es la curva IPR de la cátedra.",
    "Presión estática, presión de burbuja, índice de productividad o datos de ensayo, método elegido.",
    "IPR lineal (Darcy), Vogel (bajo el punto de burbuja), Fetkovich y combinada; y la inversa: qué Pwf corresponde al caudal objetivo.",
    "La Pwf a la profundidad de los punzados para el caudal de diseño.",
    "No habría vínculo entre el caudal deseado y la presión disponible: el diseño de la bomba partiría de un dato inventado.",
    "calculate_pwf_for_target_rate(), y las curvas IPR por método")

modulo("core/pvt.py — el fluido a condiciones de fondo",
    "El petróleo con gas disuelto cambia de volumen, viscosidad y densidad según presión y temperatura. Sin PVT, todo gradiente es erróneo.",
    "°API, gravedad del gas, GOR, temperatura, presión.",
    "Correlaciones clásicas: Standing (Pb, Rs, Bo), Beggs-Robinson (viscosidad), Dranchuk-Abou-Kassem (factor Z del gas).",
    "Propiedades del fluido a cualquier (P, T) del recorrido.",
    "Los gradientes de presión y el gas libre en la admisión quedarían mal calculados: etapas y separador mal dimensionados.",
    "Funciones por propiedad (Rs, Bo, Z, viscosidades), cada una con su correlación citada")

modulo("core/multiphase.py — del fondo a la admisión, y de la descarga a superficie",
    "Entre los punzados y la bomba (y de la bomba a superficie) fluye una mezcla de líquido y gas: la presión no cae linealmente. Es el módulo más grande porque implementa correlaciones de flujo multifásico (Hagedorn-Brown, Orkiszewski).",
    "Pwf, geometría del pozo, propiedades PVT, caudales.",
    "El recorrido de presión (traverse) por el anular hasta la profundidad de la bomba, y por el tubing hasta superficie.",
    "PIP (presión de admisión, con calculate_pip) y presión de descarga requerida (calculate_discharge_pressure).",
    "No se conocería la presión real en la admisión: sumergencia, gas libre y TDH serían suposiciones.",
    "calculate_pip(), calculate_discharge_pressure()")

modulo("core/tdh.py — la energía que debe entregar la bomba",
    "El corazón hidráulico del diseño: cuánta altura debe vencer la bomba.",
    "Profundidad de la bomba, PIP, caudal, tubería, presión requerida en boca.",
    "TDH = elevación vertical neta + fricción (Hazen-Williams) + presión en boca convertida a altura.",
    "El TDH [pies] que define el número de etapas.",
    "No habría forma de saber cuántas etapas necesita la bomba.",
    "calculate_tdh()")

modulo("core/pump_design.py — bomba y etapas",
    "Convierte TDH y caudal en una bomba concreta con su número de etapas y potencia.",
    "TDH, caudal, casing, catálogo de bombas.",
    "Filtra candidatas (diámetro y rango), interpola la curva de cada una al caudal de diseño, calcula etapas = TDH / (altura por etapa) y HP = hp/etapa × etapas × SG.",
    "La lista de diseños hidráulicos posibles (bomba, etapas, HP, eficiencia).",
    "Habría que leer las curvas del catálogo a mano para cada alternativa.",
    "design_pump_complete()")

modulo("core/electrical.py — motor, sello, cable y transformador",
    "El tren eléctrico y mecánico que acompaña a la bomba.",
    "HP requeridos, profundidad, temperatura de fondo, tensión disponible, catálogos.",
    "Selección de motor (menor HP que cubra la carga, tensión más cercana), empuje axial y sello, cable por ampacidad y temperatura con caída de tensión interpolada, tensión de superficie, kVA y transformador.",
    "El diseño eléctrico completo con sus verificaciones.",
    "La bomba quedaría sin motor dimensionado ni cable compatible: el diseño no sería instalable.",
    "electrical_design_complete(), calculate_kva(), select_transformer(), estimate_axial_thrust()")

modulo("core/gas_handling.py — el gas libre en la admisión",
    "El gas libre degrada o bloquea una bomba centrífuga. Hay que cuantificarlo y decidir el separador.",
    "PIP, PVT, GOR, corte de agua, caudal.",
    "Fracción de gas en la admisión (GIP), diseño por incrementos de presión (método de Brown), recomendación de separador desde el catálogo.",
    "La configuración de manejo de gas y sus advertencias.",
    "Diseños que fallarían en pozos con gas: el caso más común de mal desempeño de BES.",
    "complete_gas_design(), select_gas_handler()")

modulo("catalogs/loader.py y database_migration/database_loader_v3.py — la ventanilla de datos",
    "Una sola pieza sabe leer la base de datos; el resto del programa le pide 'las bombas que entran en este casing' sin saber si detrás hay JSON, Excel o SQLite.",
    "Los archivos de la base (hoy: data_excel/*.xlsx).",
    "Lecturas, uniones entre tablas (bomba+curva+housings; cable+caída de tensión), interpolaciones de catálogo y las selecciones (get_motor, get_cable, get_seal...).",
    "Objetos y diccionarios listos para los cálculos.",
    "Cada módulo leería archivos por su cuenta: cambiar el almacenamiento exigiría tocar todo el programa.",
    "ExcelCatalogManager (hereda de CatalogManager), interpolate_pump_curve(), get_pumps_by_casing(), get_motor(), get_cable(), get_seal(), select_gas_handler(), select_sensor()")

modulo("recommender/ — el comité de selección",
    "Un diseño no es 'la primera bomba que sirve' sino la mejor entre alternativas comparadas con criterios explícitos.",
    "Los datos del pozo y el catálogo.",
    "Corre el diseño hidráulico y eléctrico completo con cada candidata (pump_selector.py), puntúa eficiencia (40%), flexibilidad/cercanía al BEP (30%) y preferencia de fabricante (30%) (scoring.py), y arma el ranking con justificación (recommendation_engine.py).",
    "Las top-N recomendaciones con puntajes y razones.",
    "El usuario vería una única opción sin saber si era la mejor.",
    "generate_recommendations(), select_top_n_pumps(), overall_score()")

modulo("ui/ + app.py — la ventanilla de atención",
    "Traducir todo lo anterior a pantallas que un ingeniero usa sin leer código.",
    "Clics y valores del usuario.",
    "No calcula ingeniería: arma formularios (forms.py), muestra resultados y gráficos (results_view.py, plots.py), compara opciones (comparison_view.py) y corre sensibilidades (sensitivity_view.py).",
    "La experiencia visual y el botón de descarga del reporte Excel.",
    "El programa existiría solo para programadores.",
    "render_results() y las vistas por sección; app.py orquesta")

# ============================================================ 8. FÓRMULAS
h1("8. Fórmulas utilizadas y desarrollo matemático")
p("Cada fórmula se presenta con: qué representa, sus variables, de dónde viene y "
  "dónde vive en el programa. Unidades del proyecto: presiones en psia, "
  "temperaturas en °F, caudales en bpd, longitudes en pies, diámetros en pulgadas, "
  "potencia en HP.")

h3("8.1 IPR — capacidad de aporte (core/ipr.py)")
p("Lineal (Darcy), válida sobre el punto de burbuja: el caudal es proporcional a "
  "la depresión.")
eq("q = J · (Pr − Pwf)", "q: caudal [bpd] · J: índice de productividad [bpd/psi] · Pr: presión estática · Pwf: presión dinámica de fondo")
p("Vogel, bajo el punto de burbuja (el gas liberado curva la IPR):")
eq("q / q_max = 1 − 0.2·(Pwf/Pr) − 0.8·(Pwf/Pr)²", "Vogel (1968); usada por Brown §4.5")
p("Fetkovich, de ensayos multi-caudal:")
eq("q = C · (Pr² − Pwf²)ⁿ", "C y n se ajustan de los ensayos del pozo")
p("El programa además las invierte: dado el caudal objetivo, despeja la Pwf "
  "necesaria (calculate_pwf_for_target_rate). Ese Pwf alimenta el cálculo de PIP.")

h3("8.2 PVT — propiedades del fluido (core/pvt.py)")
tabla(["Propiedad", "Correlación", "Qué estima"], [
    ("Presión de burbuja Pb, gas disuelto Rs, volumen Bo", "Standing (1947)", "Cuánto gas está disuelto y cuánto 'se hincha' el petróleo a cada P,T."),
    ("Viscosidad del petróleo", "Beggs-Robinson (1975)", "Resistencia a fluir, muerta y con gas disuelto."),
    ("Factor Z del gas", "Dranchuk-Abou-Kassem (1975)", "Desviación del gas real respecto del ideal; da la densidad del gas."),
])
p("Por qué correlaciones y no tablas: cuando no hay PVT de laboratorio del pozo, "
  "las correlaciones son el estándar de la industria. La base de datos ya tiene la "
  "tabla fluid_samples prevista para validar contra PVT medidos cuando existan.")

h3("8.3 TDH — Altura Dinámica Total (core/tdh.py, Brown §4.5324)")
p("Qué representa: la energía por unidad de peso que la bomba debe entregar, "
  "expresada como columna de fluido [pies]. Tres sumandos:")
eq("TDH = Elevación vertical + Fricción en tubing + Altura por presión en boca")
p("Elevación vertical neta: la profundidad de la bomba menos lo que 'ayuda' la "
  "presión de admisión convertida a columna:")
eq("Elev = D_bomba − PIP · 2.31 / SG", "2.31 pies de agua ≡ 1 psi; SG: gravedad específica del líquido")
p("De dónde sale 2.31: una columna de agua dulce de 2.31 pies ejerce 1 psi "
  "(62.4 lb/ft³ · 2.31 ft / 144 in²/ft² ≈ 1 psi). Dividir por SG lo corrige para "
  "el fluido real.")
p("Fricción por Hazen-Williams (tuberías, flujo de agua/liquidos poco viscosos):")
eq("h_f = 0.2083 · (100/C)^1.852 · q_gpm^1.852 / d^4.8655   [pies / 100 pies] · L/100",
   "C: coeficiente de rugosidad (~120 tubería usada) · q en gpm · d: diámetro interno [in] · L: longitud [pies]")
p("Presión requerida en boca, convertida a altura:")
eq("h_whp = P_wh · 2.31 / SG")
p("Validación: el ejemplo de fricción de Brown (tubería 5\", 10.000 bpd) da "
  "≈18.5 pies/1000 pies y es un test automático del proyecto.")

h3("8.4 Etapas y potencia (core/pump_design.py)")
eq("N_etapas = TDH / (altura por etapa al caudal de diseño)")
eq("HP = (hp por etapa) · N_etapas · SG",
   "Las curvas de catálogo están medidas con agua (SG=1): la potencia escala con la densidad del fluido real")
p("La 'altura por etapa' y el 'hp por etapa' salen de interpolar la curva "
  "digitalizada de la bomba al caudal de diseño (capítulo 10). El número de etapas "
  "se materializa luego con los housings disponibles del catálogo.")

h3("8.5 Diseño eléctrico (core/electrical.py)")
p("Caída de tensión en el cable — el conductor tiene resistencia y se calienta:")
eq("ΔV = v(T) · I · L / 1000",
   "v(T): caída [V/A/1000 ft] del catálogo, interpolada a la temperatura del conductor · I: corriente de placa · L: profundidad de la bomba")
p("Tensión en superficie (agrega la pérdida típica del transformador, 2.5%):")
eq("V_s = (V_motor + ΔV_cable) · 1.025")
p("Potencia aparente y transformador (alimentación trifásica):")
eq("kVA = V_s · I · √3 / 1000", "√3 = 1.732: idéntica a la fórmula de la cátedra")
p("Selección: el menor tamaño estándar de transformador con kVA ≥ demanda "
  "(select_transformer). Nunca subdimensionar: la tabla transformers de la base "
  "contiene los tamaños estándar (25–300 kVA), extraídos del código a la base "
  "durante la fase 4.")
p("Empuje axial y sello: el empuje hidráulico se estima con la presión "
  "diferencial de la bomba actuando sobre el área del eje, con margen de diseño "
  "1.2 (estimate_axial_thrust); el sello se elige por capacidad de cojinete ≥ "
  "empuje, temperatura y compatibilidad de serie, prefiriendo laberinto por "
  "economía (get_seal) — el mismo criterio de la cátedra.")

h3("8.6 La identidad hidráulica (control de calidad de curvas)")
p("Deducción completa, porque es el candado físico de la digitalización. La "
  "potencia hidráulica es peso por unidad de tiempo por altura:")
eq("P_hidr = γ · Q · H")
p("Pasando a unidades de campo: Q [bpd] · 5.615 ft³/bbl / 86.400 s/día; "
  "γ = 62.4 · SG lb/ft³; 1 HP = 550 lb·ft/s:")
eq("HP_hidr = Q · H · SG · 62.4 · 5.615 / (86.400 · 550) = Q · H · SG / 135.773")
eq("η = HP_hidr / HP_freno = Q · H · SG / (135.773 · HP)")
p("En cada punto digitalizado, la eficiencia leída del gráfico debe coincidir con "
  "la calculada con esta identidad. Si no cierra, la lectura fue mala y la bomba "
  "se rechaza para revisión. Así se validaron las 37 bombas Alkhorayef.")

# ============================================================ 9-10. CURVAS
h1("9. Curvas BES: qué son y cómo viven en el sistema")
p("Una curva de rendimiento describe, por etapa y a 60 Hz con agua: la altura que "
  "entrega (decreciente con el caudal), la potencia que consume y la eficiencia "
  "(con un máximo en el BEP, Best Efficiency Point). El fabricante la publica como "
  "gráfico; el programa la necesita como números.")
img_curva = ROOT / "database_migration" / "_curva_wa550.png"
if img_curva.exists():
    doc.add_picture(str(img_curva), width=Inches(6.0))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Figura 2 — Curva de catálogo (Alkhorayef WA-550) tal como la procesa el extractor")
    cr.italic = True; cr.font.size = Pt(9)
p("Recorrido completo de una curva dentro del sistema:")
diagrama([
    "  Gráfico del catálogo (PDF)",
    "        │  extractor: OCR de ejes + separación por color",
    "        ▼",
    "  ~10 puntos (Q, H/etapa, HP/etapa, η)  +  control de calidad físico",
    "        │  se guardan como filas",
    "        ▼",
    "  Tabla pump_curves (pumps.xlsx) — clave: bomba + caudal",
    "        │  la ventanilla las une con su bomba",
    "        ▼",
    "  Interpolación lineal al caudal de diseño (interpolate_pump_curve)",
    "        │",
    "        ▼",
    "  Etapas, potencia y eficiencia del diseño  →  gráficos en ui/plots.py",
])

h1("10. Regresiones e interpolaciones")
p("Decisión de diseño: las curvas se guardan como PUNTOS del catálogo y el "
  "programa interpola linealmente entre ellos. No se ajustan polinomios "
  "(regresión) para las curvas de bombas. Por qué: cada punto guardado es "
  "verificable contra el catálogo impreso (transparencia), y un polinomio puede "
  "oscilar entre puntos e inventar valores sin respaldo físico.")
eq("y = y₁ + (y₂ − y₁) · (x − x₁) / (x₁₂ distancia)",
   "Interpolación lineal: regla de tres en el tramo entre dos puntos consecutivos")
p("Dónde SÍ se usan regresiones/ajustes en el proyecto:")
bullets([
    "Fuera de rango, el programa NO extrapola: avisa que el caudal está fuera de la curva (bounds_error), en lugar de inventar.",
    "La caída de tensión de cables se interpola linealmente en temperatura entre los valores de catálogo.",
    "En la digitalización de curvas se usó un ajuste lineal robusto (tipo RANSAC) SOLO para calibrar los ejes de los gráficos (relación píxel→unidad), nunca para modificar los datos de las curvas.",
    "Las correlaciones PVT (Standing, Beggs-Robinson) son, en sí mismas, regresiones publicadas de la industria: se usan citadas, no re-ajustadas.",
])

# ============================================================ 11. SELECCIÓN
h1("11. Metodología de selección de equipos")
p("Sigue el procedimiento de la cátedra y de Brown, automatizado:")
tabla(["Paso (cátedra)", "Dónde vive", "Criterio"], [
    ("1. Datos del pozo", "ui/forms.py + core/models.py", "Formularios validados."),
    ("2. Bomba y etapas", "pump_design.py + recommender", "Candidatas que entran en el casing y cubren el caudal; etapas por TDH; se comparan 2-3+ alternativas."),
    ("3. Verificaciones mecánicas", "electrical.py (empuje/eje) — MaxP y refrigeración: previstas (datos ya en la base)", "Empuje ≤ capacidad del sello; límites de eje del catálogo."),
    ("4. Motor", "electrical.py + get_motor()", "Menor HP ≥ carga; tensión más alta disponible para bajar corriente."),
    ("5. Sello", "get_seal()", "Empuje, temperatura, serie compatible; laberinto preferido."),
    ("6. Controlador", "Tabla switchboards (estructura lista)", "Por V/A/HP máximos (pendiente de datos Wood Group)."),
    ("7. Cable", "get_cable()", "Ampacidad y temperatura; menor caída de tensión."),
    ("8. Transformador", "calculate_kva() + select_transformer()", "kVA = V·I·√3/1000; menor tamaño ≥ demanda."),
])
p("Puntuación de alternativas (recommender/scoring.py): eficiencia 40% + "
  "flexibilidad 30% (cuán cerca del BEP opera: una bomba lejos de su BEP vibra, "
  "desgasta y desperdicia energía) + preferencia de fabricante 30% (contratos "
  "reales condicionan la marca). Los pesos son explícitos y defendibles; el "
  "criterio económico se incorporará cuando exista la tabla equipment_costs.")

# ============================================================ 12. STREAMLIT
h1("12. Desarrollo de la interfaz (Streamlit)")
p("Qué es Streamlit: una biblioteca que convierte un programa Python en una "
  "aplicación web local con formularios, pestañas y gráficos, sin desarrollar una "
  "página web desde cero. Se eligió porque: (1) el esfuerzo va a la ingeniería y "
  "no al diseño web; (2) es el estándar de facto para herramientas técnicas en "
  "Python; (3) cualquier ingeniero la ejecuta con un comando:")
diagrama(["  python -m streamlit run app.py"])
p("Secciones de la aplicación: carga de datos del pozo → Diseño BES (ranking de "
  "recomendaciones con pestañas por alternativa) → Comparación de opciones → "
  "Análisis de sensibilidad → descarga del reporte Excel. La interfaz NO calcula "
  "nada: muestra lo que los módulos de core/ y recommender/ calcularon — por eso "
  "un error visual nunca puede alterar un resultado de ingeniería.")

# ============================================================ 13. VALIDACIÓN
h1("13. Validación y auditoría")
p("Tres capas independientes, todas automatizadas y repetibles:")
tabla(["Capa", "Herramienta", "Qué garantiza"], [
    ("1. Validación contra el libro", "tests/ (663 ensayos, pytest)", "Los ejemplos resueltos de Brown se reproducen: Ej. 2A → TDH 5.830 ft, 254 etapas, bomba D-40. Si un cambio rompe un cálculo, el ensayo falla al instante."),
    ("2. Integridad de la base", "check_integrity.py", "Claves únicas, ninguna referencia sin destino, rangos físicos (0<η<1, min<BEP<max), relaciones 1:1 completas."),
    ("3. Equivalencia del origen de datos", "verify_database_v31.py", "Las 23 bombas históricas idénticas al origen JSON dentro del catálogo ampliado; las selecciones legadas no cambian."),
])
p("Auditoría externa del modelo de datos: antes de congelar el diseño se hizo una "
  "revisión crítica (documento AUDITORIA_BASE_DE_DATOS.md) que encontró y corrigió "
  "una violación de tercera forma normal (caída de tensión duplicada), una "
  "referencia huérfana (serie 513) y remodeló los casos de campo para representar "
  "instalaciones reales en tándem. Que la auditoría haya encontrado defectos y "
  "queden documentados es un punto A FAVOR en una defensa: demuestra método.")
p("Cifra final tras integrar la base a la aplicación: 663 de 663 ensayos en verde.", bold=True)

# ============================================================ 14. AMPLIAR
h1("14. Cómo agregar nuevos fabricantes")
p("Regla de oro: agregar un fabricante es agregar FILAS, nunca código.")
bullets([
    "Paso 1 — Alta del fabricante: una fila en manufacturers.xlsx.",
    "Paso 2 — Fuente: una fila en data_sources.xlsx con el catálogo citado (ej. WG-01).",
    "Paso 3 — Equipos: filas en pumps/motors/cables/seals... usando el mismo manufacturer_id. Si aparece una serie nueva, alta en equipment_series.xlsx.",
    "Paso 4 — Curvas de bombas: filas en pump_curves (10 puntos por bomba). Si el catálogo es PDF con gráficos, el extractor de Alkhorayef sirve de plantilla.",
    "Paso 5 — Verificar: correr check_integrity.py (debe dar 0 errores) y abrir la aplicación: las candidatas nuevas aparecen solas.",
])
p("Demostración ya realizada: Alkhorayef entró con 37 bombas, 376 housings y 621 "
  "puntos de curva sin modificar una línea de la lógica de la aplicación.")

# ============================================================ 15. MANTENER
h1("15. Cómo mantener el sistema")
tabla(["Situación", "Qué hacer"], [
    ("Editaste un Excel de la base a mano", "Correr check_integrity.py. Si imprime 'INTEGRIDAD OK', la edición es segura."),
    ("Cambiaste código de cálculo", "Correr pytest. Los 663 ensayos deben quedar en verde; si uno falla, te dice exactamente qué ejemplo del libro dejó de reproducirse."),
    ("Querés regenerar la base desde cero", "build_database_v3.py → import_alkhorayef.py → extract_curves_alkhorayef.py → check_integrity.py (orden documentado en database_migration/README.md)."),
    ("Querés volver al origen de datos anterior", "En app.py, función _load_catalog: descomentar la línea del CatalogManager JSON (está comentada como respaldo)."),
    ("Vas a migrar a SQLite", "El esquema ya es relacional: cada hoja se convierte en CREATE TABLE con las PK/FK documentadas en los README de cada Excel y los índices listados en DISENO_BASE_DE_DATOS.md."),
    ("Duda sobre un dato", "Toda fila tiene source_id: buscarlo en data_sources.xlsx y abrir el PDF citado en TESIS/CATALOGOS."),
])

# ============================================================ 16. DEFENSA
h1("16. Cómo defender académicamente el proyecto")
p("Las preguntas probables y su respuesta corta:")
tabla(["Pregunta esperable", "Respuesta"], [
    ("¿Por qué confiar en los resultados?", "Porque reproducen los ejemplos resueltos de Brown, y esa validación corre automáticamente (663 ensayos) tras cada cambio."),
    ("¿Por qué Excel como base de datos?", "Editable y auditable por ingenieros sin programar; misma estructura lógica que SQLite: la migración futura cambia el motor, no el diseño."),
    ("¿Por qué separar datos y lógica?", "Los catálogos cambian, el método no. Se demostró: un fabricante completo entró sin tocar código."),
    ("¿Por qué digitalizar catálogos?", "Un gráfico no se puede consultar ni comparar; una tabla sí. Y la digitalización se validó con física (identidad hidráulica, BEP), no a ojo."),
    ("¿Por qué interpolar y no ajustar polinomios?", "Transparencia: cada punto es verificable contra el catálogo impreso; un polinomio puede inventar valores entre puntos."),
    ("¿Por qué Streamlit?", "El esfuerzo del proyecto es de ingeniería, no de desarrollo web; Streamlit da la interfaz con costo mínimo y sin tocar los cálculos."),
    ("¿Esto sigue lo que enseña la cátedra?", "Sí: el mapa paso a paso está en METODOLOGIA_CATEDRA_VS_APLICACION.md; los pasos 1, 2, 7 y 8 completos, y las verificaciones mecánicas/térmicas adicionales ya tienen sus datos en la base."),
    ("¿Qué encontró la auditoría?", "Defectos reales (3FN en cables, serie huérfana), corregidos y documentados: evidencia de método, no de improvisación."),
    ("¿Cómo crece a futuro?", "Nuevos fabricantes = filas; nuevas verificaciones = funciones que leen columnas ya creadas; SQLite = mismo esquema con motor formal."),
])
p("Estrategia de defensa sugerida: abrir con la separación método/datos, mostrar "
  "el ERD como plano de la base, correr en vivo pytest y check_integrity (30 "
  "segundos, todo en verde) y cerrar con la demostración Alkhorayef: 'agregué un "
  "fabricante entero sin programar'.")

# ============================================================ CIERRE
h1("17. Referencias del proyecto")
p("Documentos que amplían cada capítulo (en database_migration/ y en "
  "TESIS/Defensa BES Designer/ en versión Word):")
archivo("DISENO_BASE_DE_DATOS.md", "el diseño completo del modelo de datos, tabla por tabla.")
archivo("AUDITORIA_BASE_DE_DATOS.md", "la auditoría técnica externa y sus hallazgos.")
archivo("INVENTARIO_BASES_DE_DATOS.md", "el inventario de bases explicado para la defensa.")
archivo("METODOLOGIA_CATEDRA_VS_APLICACION.md", "el mapa paso a paso contra la metodología de la cátedra.")
archivo("INFORME_COBERTURA_CATALOGOS.md", "qué catálogos de TESIS están incorporados y cuáles faltan.")
archivo("erd.png / erd.svg / erd.mermaid", "el diagrama entidad-relación en tres formatos.")
archivo("docs/METHODOLOGY.md (proyecto)", "la metodología de cálculo del motor original.")
p("")
p("Bibliografía principal: Brown, K.E. (1980), The Technology of Artificial Lift "
  "Methods, Vol. 2B, Cap. 4.5 · Standing (1947) · Vogel (1968) · Beggs & Robinson "
  "(1975) · Dranchuk & Abou-Kassem (1975) · API RP 11S · Takacs (2018), Electrical "
  "Submersible Pumps Manual · Catálogos de fabricantes citados en data_sources.xlsx.")

doc.save(OUT)
print(f"OK — {OUT} ({OUT.stat().st_size} bytes)")
